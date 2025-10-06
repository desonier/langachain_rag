import os
import streamlit as st
from dotenv import load_dotenv
import base64
import docx
import json
import time
import re
import random
import PyPDF2
from io import BytesIO
from openai import AzureOpenAI
from azure.storage.blob import BlobServiceClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents import SearchClient
from azure.search.documents.models import VectorizedQuery
from azure.identity import DefaultAzureCredential
import tiktoken
import numpy as np 

class BlobUploader:
    def __init__(self, connection_string, container_name):
        self.client = BlobServiceClient.from_connection_string(connection_string)
        self.container = self.client.get_container_client(container_name)

    def upload_file(self, local_path, blob_name=None):
        if not blob_name:
            blob_name = os.path.basename(local_path)
        blob_client = self.container.get_blob_client(blob_name)
        with open(local_path, "rb") as data:
            blob_client.upload_blob(data, overwrite=True)
        print(f"Uploaded {local_path} to blob storage as {blob_name}")
        return blob_name, blob_client

    def extract_text(self, blob_name, blob_client):
        data = blob_client.download_blob().readall()
        if blob_name.lower().endswith('.txt'):
            return data.decode("utf-8", errors="ignore")
        elif blob_name.lower().endswith('.docx'):
            doc = docx.Document(BytesIO(data))
            return "\n".join([para.text for para in doc.paragraphs])
        elif blob_name.lower().endswith('.pdf'):
            pdf = PyPDF2.PdfReader(BytesIO(data))
            return "\n".join([page.extract_text() or "" for page in pdf.pages])
        else:
            return ""

    def get_blob_metadata(self, blob_client):
        props = blob_client.get_blob_properties()
        safe_key = re.sub(r'[^a-zA-Z0-9_\-=]', '_', blob_client.blob_name)
        return {
            "Resume_ID": safe_key,
            "metadata_storage_name": blob_client.blob_name,
            "metadata_storage_size": props.size,
            "metadata_storage_last_modified": props.last_modified.isoformat()
        }

class OpenAIHelper:
    def __init__(self, endpoint, api_key, api_version, embedding_model, chat_model):
        self.client = AzureOpenAI(
            azure_endpoint=endpoint,
            api_key=api_key,
            api_version=api_version
        )
        self.embedding_model = embedding_model
        self.chat_model = chat_model
        self.tokenizer = tiktoken.encoding_for_model(embedding_model)

    def embed_text(self, text, max_tokens=8192):  # <-- Change max_tokens to 8192
        tokens = self.tokenizer.encode(text)
        if len(tokens) > max_tokens:
            print(f"⚠️ Text exceeds {max_tokens} tokens ({len(tokens)}). Chunking and averaging embeddings.")
            chunk_embeddings = []
            for i in range(0, len(tokens), max_tokens):
                chunk = self.tokenizer.decode(tokens[i:i+max_tokens])
                response = self.client.embeddings.create(model=self.embedding_model, input=[chunk])
                chunk_embeddings.append(response.data[0].embedding)
            # Average embeddings to get a single vector of correct length
            avg_embedding = np.mean(chunk_embeddings, axis=0)
            return avg_embedding.tolist()
        else:
            response = self.client.embeddings.create(model=self.embedding_model, input=[text])
            return response.data[0].embedding

    
    def score_resume(self, user_prompt, resume_text):
        relevance_prompt = [
        {
            "role": "system",
            "content": (
                "You are a resume analyst and search assistant. "
                "Your role is to retrieve, evaluate, and rank resumes based on how well they match a given job description. "
                "You must provide accurate, transparent, and structured feedback for each resume returned with a pros and cons. "
                "For each selected resume, provide a high level summary at the top why the resume was selected and follow the template below.\n"
                "Summary:\n"
                "Provide a concise overview of the candidate’s overall fit for the role.\n"
                "Pros:\n"
                "List strengths under the following categories:\n"
                "•  Educational Qualification: Degrees, institutions, relevance to role\n"
                "•  Experience: Years of experience, seniority, domain relevance\n"
                "•  Technical Skills: Specific tools, platforms, languages, proficiencies\n"
                "•  Certifications: Industry-recognized credentials\n"
                "•  Security Clearance: If applicable\n"
                "•  Leadership: Team management, project ownership, strategic roles\n"
                "Cons:\n"
                "List gaps or limitations under the following categories:\n"
                "•  Skill Gaps: Missing or underdeveloped technical proficiencies\n"
                "•  Domain Fit: Limited experience in target industry or role type\n"
                "•  Tool Specificity: Lack of detail in tools or scripting languages\n"
                "•  Role Alignment: Misalignment with job title or responsibilities\n"
                "Conclusion:\n"
                "Summarize the candidate’s viability for the role. Include a match rating from 1 to 10, and explain whether they are a strong, moderate, or weak fit based on the job description."
            )
        },
        {
            "role": "user",
            "content": f"Job Description:\n{user_prompt}\n\nResume:\n{resume_text}"
        }
    ]
        response = self.client.chat.completions.create(
            model=self.chat_model,
            messages=relevance_prompt,
            temperature=0.8,
            max_tokens=500
        )
        return response.choices[0].message.content

    def extract_fields_with_gpt(self, text):
        prompt = """
Extract the following fields from this resume text and return them as JSON.
Please format Certifications as a comma-separated string.

Fields:
- Name
- Location
- Phone
- Email
- Clearance Level
- Education
- Certifications
- Skills

Resume:
""" + text
        try:
            response = self.client.chat.completions.create(
                model=self.chat_model,
                messages=[{"role": "user", "content": prompt}]
            )
            content = response.choices[0].message.content.strip()
            json_start = content.find('{')
            json_end = content.rfind('}') + 1
            json_text = content[json_start:json_end]
            fields = json.loads(json_text)
        except Exception as e:
            print(f"⚠️ Error parsing GPT response: {e}")
            print(f"Raw GPT output:\n{content}")
            fields = {}
        return fields

    @staticmethod
    def flatten_field(value):
        if isinstance(value, dict):
            return ", ".join(f"{k}: {v}" for k, v in value.items())
        elif isinstance(value, list):
            return ", ".join(str(v) for v in value)
        elif isinstance(value, (str, int, float)):
            return str(value)
        else:
            return None

    @staticmethod
    def flatten_field(value):
        if isinstance(value, dict):
            return ", ".join(f"{k}: {v}" for k, v in value.items())
        elif isinstance(value, list):
            return ", ".join(str(v) for v in value)
        elif isinstance(value, (str, int, float)):
            return str(value)
        else:
            return None

class AzureSearchIndexer:
    def __init__(self, endpoint, api_key, index_name):
        self.client = SearchClient(endpoint, index_name, AzureKeyCredential(api_key))

    def build_document(self, blob_client, embedding, full_text, meta, fields):
        # Chunk resume text for storage
        tokenizer = tiktoken.encoding_for_model("text-embedding-ada-002")
        tokens = tokenizer.encode(full_text)
        chunk_size = 8192  # <-- Change chunk_size to 8192
        chunks = [tokenizer.decode(tokens[i:i+chunk_size]) for i in range(0, len(tokens), chunk_size)]
        doc = {
            "@search.action": "upload",
            "content": full_text,
            "contentVector": embedding,
            "resumeText": full_text,
            "chunks": chunks,  # <-- Store chunked text
            **meta,
            "Resume_Date": None,
            "Name": fields.get("Name"),
            "Location": fields.get("Location"),
            "Phone": fields.get("Phone"),
            "Email": fields.get("Email"),
            "Clearance_Level": fields.get("Clearance Level"),
            "Education": OpenAIHelper.flatten_field(fields.get("Education")),
            "Certifications": OpenAIHelper.flatten_field(fields.get("Certifications")),
            "Skills": OpenAIHelper.flatten_field(fields.get("Skills")),
            "Certifications_Clearances": None,
            "Professional_Experience": None,
            "Work_Experience": None,
            "Related_Experience": None,
            "Source": None,
            "metadata_author": None,
            "metadata_title": None,
            "metadata_subject": None,
            "metadata_keywords": None,
            "metadata_comments": None,
        }
        return doc

    def upload_document(self, doc):
        result = self.client.upload_documents(documents=[doc])
        # print("Indexed document in Azure Cognitive Search.")

    def validate_index(self, resume_id):
        for _ in range(10):
            results = self.client.search(search_text=resume_id, select=["Resume_ID"])
            for doc in results:
                if doc.get("Resume_ID") == resume_id:
                    # print(f"Blob {resume_id} successfully indexed.")
                    return True
            time.sleep(2)
        # print(f"Blob {resume_id} was NOT found in the index after waiting.")
        return False

class ResumeIngestionPipeline:
    def __init__(self):
        load_dotenv()
        self.blob_uploader = BlobUploader(
            os.getenv("AZURE_STORAGE_CONNECTION_STRING"),
            os.getenv("CONTAINER_NAME")
        )
        self.openai_helper = OpenAIHelper(
            os.getenv("AZURE_OPENAI_ENDPOINT"),
            os.getenv("AZURE_OPENAI_KEY"),
            os.getenv("AZURE_OPENAI_API_VERSION"),
            os.getenv("EMBEDDING_MODEL"),
            os.getenv("OPENAI_MODEL", "gpt-35-turbo")
        )
        self.indexer = AzureSearchIndexer(
            os.getenv("AZURE_AI_SEARCH_ENDPOINT"),
            os.getenv("AZURE_AI_SEARCH_KEY"),
            os.getenv("AZURE_AI_SEARCH_INDEX")
        )

    def process_resume_with_retry(self, local_path, max_retries=5):
        for attempt in range(max_retries):
            try:
                return self.process_resume(local_path)
            except Exception as e:
                if "429" in str(e) or "rate limit" in str(e).lower():
                    wait_time = (2 ** attempt) + random.uniform(0, 1)
                    print(f"Rate limited (429). Retrying in {wait_time:.2f} seconds...")
                    time.sleep(wait_time)
                else:
                    raise
        raise Exception("Max retries exceeded due to rate limiting.")

    def process_resume(self, local_path):
        blob_name, blob_client = self.blob_uploader.upload_file(local_path)
        text = self.blob_uploader.extract_text(blob_name, blob_client)
        if not text.strip():
            raise Exception(f"Blob {blob_name} is empty or unreadable.")
        embedding = self.openai_helper.embed_text(text)
        meta = self.blob_uploader.get_blob_metadata(blob_client)
        fields = self.openai_helper.extract_fields_with_gpt(text)
        doc = self.indexer.build_document(blob_client, embedding, text, meta, fields)
        self.indexer.upload_document(doc)
        resume_id = meta["Resume_ID"]
        success = self.indexer.validate_index(resume_id)
        return success

class AzureBlobHelper:
    def __init__(self, connection_string, container_name):
        self.client = BlobServiceClient.from_connection_string(connection_string)
        self.container = self.client.get_container_client(container_name)

    def upload_file(self, file):
        blob_client = self.container.get_blob_client(file.name)
        blob_client.upload_blob(file, overwrite=True)

    def list_blobs(self):
        return [blob.name for blob in self.container.list_blobs()]

class AzureSearchHelper:
    def __init__(self, endpoint, index_name, api_key, semantic_config):
        self.client = SearchClient(endpoint, index_name, AzureKeyCredential(api_key))
        self.semantic_config = semantic_config

    def search(self, user_prompt, prompt_vector, mode, top=5):
        base_args = dict(
            select=[
                "Resume_ID", "Resume_Date", "Name", "Location", "Phone", "Email", "Clearance_Level",
                "Education", "Certifications", "Certifications_Clearances", "Professional_Experience",
                "Work_Experience", "Related_Experience", "Skills", "Source", "content", "resumeText",
                "metadata_storage_name", "metadata_storage_size", "metadata_storage_last_modified",
                "metadata_author", "metadata_title", "metadata_subject", "metadata_keywords", "metadata_comments"
            ],
            top=top,
            include_total_count=True
        )
        if mode == "Hybrid (Semantic + Vector)":
            return self.client.search(
                search_text=user_prompt,
                vector_queries=[VectorizedQuery(
                    vector=prompt_vector,
                    k_nearest_neighbors=top,
                    fields="contentVector",
                    kind="vector",
                    exhaustive=True
                )],
                query_type="semantic",
                semantic_configuration_name=self.semantic_config,
                highlight_pre_tag="<mark>",
                highlight_post_tag="</mark>",
                **base_args
            )
        elif mode == "Semantic Only":
            return self.client.search(
                search_text=user_prompt,
                query_type="semantic",
                semantic_configuration_name=self.semantic_config,
                highlight_pre_tag="<mark>",
                highlight_post_tag="</mark>",
                **base_args
            )
        elif mode == "Vector Only":
            return self.client.search(
                search_text="",
                vector_queries=[VectorizedQuery(
                    vector=prompt_vector,
                    k_nearest_neighbors=top,
                    fields="contentVector",
                    kind="vector",
                    exhaustive=True
                )],
                **base_args
            )
def format_gpt_response(response_text):
    # Replace bullets with Markdown dashes
    lines = response_text.splitlines()
    formatted_lines = []
    for line in lines:
        # Convert bullets to Markdown
        if line.strip().startswith("•"):
            formatted_lines.append(f"- {line.strip()[1:].strip()}")
        else:
            formatted_lines.append(line)
    return "\n".join(formatted_lines)


class ResumeMatcherApp:
    def __init__(self):
        load_dotenv()
        self.search_helper = AzureSearchHelper(
            os.getenv("AZURE_AI_SEARCH_ENDPOINT"),
            os.getenv("AZURE_AI_SEARCH_INDEX"),
            os.getenv("AZURE_AI_SEARCH_KEY"),
            "default"
        )
        self.blob_helper = AzureBlobHelper(
            os.getenv("AZURE_STORAGE_CONNECTION_STRING"),
            os.getenv("CONTAINER_NAME")
        )
        self.openai_helper = OpenAIHelper(
            os.getenv("AZURE_OPENAI_ENDPOINT"),
            os.getenv("AZURE_OPENAI_KEY"),
            os.getenv("AZURE_OPENAI_API_VERSION"),
            os.getenv("EMBEDDING_MODEL"),
            os.getenv("AZURE_OPENAI_CHATGPT_DEPLOYMENT")
        )
        self.storage_account = os.getenv("AZURE_STORAGE_ACCOUNT_NAME")
        self.container_name = os.getenv("CONTAINER_NAME")

    def run(self):
        st.set_page_config(page_title="ResumeAI Matcher", layout="wide")
        self.sidebar()
        self.main_ui()


    def sidebar(self):
        import sys
        import os
        st.sidebar.header("🛠️ Admin Tools")
        uploaded_files = st.sidebar.file_uploader("Upload Resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True)
        if uploaded_files:
            pipeline = ResumeIngestionPipeline()
            st.sidebar.success(f"{len(uploaded_files)} resume(s) uploaded.")
            for file in uploaded_files:
                st.sidebar.write(f"- {file.name}")
                try:
                    # Save to a temp directory
                    temp_dir = "temp_uploaded_resume"
                    os.makedirs(temp_dir, exist_ok=True)
                    temp_file_path = os.path.join(temp_dir, file.name)
                    with open(temp_file_path, "wb") as f:
                        f.write(file.getbuffer())
                    # Pass the file path to the pipeline
                    #success = pipeline.process_resume(temp_file_path)
                    success = pipeline.process_resume_with_retry(temp_file_path)
                    if success:
                        st.sidebar.info(f"Uploaded and indexed {file.name}.")
                    else:
                        st.sidebar.error(f"Failed to index {file.name}.")
                    # Optionally, clean up temp file
                    try:
                        os.remove(temp_file_path)
                    except Exception:
                        pass
                except Exception as e:
                    st.sidebar.error(f"Failed to upload and index {file.name}: {str(e)}")
        

    def main_ui(self):
        st.markdown("<h1>SageCor Resume Intelligence.  Ver. 0.0.9 </h1>", unsafe_allow_html=True)
        user_prompt = st.text_area("Job Description", height=200)
        # --- Add file upload for job description ---
        uploaded_description_file = st.file_uploader("Or upload a job description file", type=["txt", "pdf", "docx"])
        if uploaded_description_file:
            file_extension = uploaded_description_file.name.split('.')[-1]
            if file_extension == "txt":
                job_description = uploaded_description_file.read().decode("utf-8")
            elif file_extension == "pdf":
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(uploaded_description_file)
                job_description = ""
                for page in pdf_reader.pages:
                    job_description += page.extract_text()
            elif file_extension == "docx":
                import docx
                doc = docx.Document(uploaded_description_file)
                job_description = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            else:
                job_description = ""
            user_prompt = job_description
            st.session_state['uploaded_job_description'] = job_description

        search_mode = st.selectbox(
            "Select search mode:",
            ("Hybrid (Semantic + Vector)", "Semantic Only", "Vector Only"),
            index=0
        )
        if st.button("Find Matches"):
            if user_prompt:
                with st.spinner("Embedding prompt and searching resumes..."):
                    prompt_vector = self.openai_helper.embed_text(user_prompt)
                    results = self.search_helper.search(user_prompt, prompt_vector, search_mode)
                    st.session_state.results = list(results)
                    self.display_results(user_prompt)
            else:
                st.warning("Please enter a job description or upload a job description file.")

    def display_results(self, user_prompt):
        if not st.session_state.results:
            st.warning("No resumes matched your query. Try a different prompt or check your index.")
            return
        st.subheader("📄 Matching Resumes")
        st.session_state.export_lines = []
        for i, doc in enumerate(st.session_state.results):
            blob_name = doc.get("metadata_storage_name")
            resume_link = f"https://{self.storage_account}.blob.core.windows.net/{self.container_name}/{blob_name}"
            st.session_state.export_lines.append(f"- [{blob_name}]({resume_link})")
            st.markdown(f"### 🧾 Resume {i+1}")
            st.markdown(f"- [{blob_name}]({resume_link})")
            st.markdown(f"""
                **Name:** {doc.get('Name', 'N/A')}  
                **Location:** {doc.get('Location', 'N/A')}  
                **Email:** {doc.get('Email', 'N/A')}
                """)
            captions = doc.get("@search.captions")
            if captions:
                st.markdown("**Semantic Caption:**")
                for cap in captions:
                    st.markdown(f"> {cap['text']}")
            highlights = doc.get("@search.highlights")
            if highlights:
                st.markdown("**Highlights:**")
                for field, snippets in highlights.items():
                    for snippet in snippets:
                        st.markdown(f"- {snippet}")
            resume_text = doc.get("resumeText", "")
            if resume_text.strip():
                with st.spinner(f"Scoring resume {i+1}..."):
                    # print("User Prompt:", user_prompt)
                    summary = self.openai_helper.score_resume(user_prompt, resume_text)
                    # print("Relevance Summary:", summary)
                with st.expander(f"📊 Resume {i+1} Relevance Score & Justification"):
                    formatted = format_gpt_response(summary)
                    st.markdown(formatted)
        st.subheader("📤 Export Resume Links")
        export_text = "\n".join(st.session_state.export_lines)
        st.text_area("Resume Links (Markdown format)", value=export_text, height=200, key="resume_links_export")

# --- Run App ---
if __name__ == "__main__":
    app = ResumeMatcherApp()
    app.run()